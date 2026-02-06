"""
Document Exporter
Exports OCR results to various formats: TXT, JSON, CSV, DOCX, PDF
"""
import json
import csv
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import asdict
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cv2
import numpy as np

from document_ai.core.ocr.pipeline import OCROutput
from document_ai.core.ocr.detector.craft_detector import TextBox
from document_ai.document.loader.document_loader import Document
from document_ai.exceptions import ExportException
from document_ai.config import ExportFormat
from document_ai.utils.logging import get_logger

logger = get_logger(__name__)


class DocumentExporter:
    """
    Export OCR results to various formats
    Handles formatting, styling, and file generation
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize exporter
        
        Args:
            output_dir: Default output directory
        """
        from document_ai.config import DEFAULT_CONFIG
        self.output_dir = output_dir or DEFAULT_CONFIG.storage.output_path
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export(
        self,
        ocr_results: List[OCROutput],
        output_path: Path,
        export_format: ExportFormat,
        document: Optional[Document] = None,
    ) -> Path:
        """
        Export OCR results in specified format
        
        Args:
            ocr_results: List of OCR outputs (one per page)
            output_path: Output file path
            export_format: Desired export format
            document: Original document (needed for some exports)
        
        Returns:
            Path to exported file
        
        Raises:
            ExportException: If export fails
        """
        try:
            logger.info(f"Exporting to {export_format.value} format: {output_path}")
            
            if export_format == ExportFormat.TXT:
                return self._export_txt(ocr_results, output_path)
            elif export_format == ExportFormat.JSON:
                return self._export_json(ocr_results, output_path)
            elif export_format == ExportFormat.CSV:
                return self._export_csv(ocr_results, output_path)
            elif export_format == ExportFormat.DOCX:
                return self._export_docx(ocr_results, output_path)
            elif export_format == ExportFormat.PDF:
                return self._export_pdf(ocr_results, output_path, document)
            elif export_format == ExportFormat.SEARCHABLE_PDF:
                return self._export_searchable_pdf(ocr_results, output_path, document)
            else:
                raise ExportException(f"Unsupported export format: {export_format}")
        
        except Exception as e:
            raise ExportException(f"Export failed: {str(e)}")
    
    def _export_txt(self, ocr_results: List[OCROutput], output_path: Path) -> Path:
        """Export as plain text"""
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, result in enumerate(ocr_results):
                if len(ocr_results) > 1:
                    f.write(f"--- Page {i+1} ---\n")
                f.write(result.full_text)
                f.write("\n\n")
        
        logger.info(f"Exported {len(ocr_results)} pages to TXT")
        return output_path
    
    def _export_json(self, ocr_results: List[OCROutput], output_path: Path) -> Path:
        """Export as structured JSON"""
        data = {
            "total_pages": len(ocr_results),
            "pages": []
        }
        
        for i, result in enumerate(ocr_results):
            page_data = {
                "page_number": i + 1,
                "full_text": result.full_text,
                "confidence": result.confidence,
                "num_regions": result.num_detected_regions,
                "text_regions": []
            }
            
            for box in result.text_boxes:
                x, y, w, h = box.bbox
                region_data = {
                    "text": box.text,
                    "confidence": box.confidence,
                    "bounding_box": {
                        "x": int(x),
                        "y": int(y),
                        "width": int(w),
                        "height": int(h),
                    },
                    "points": box.points.tolist(),
                }
                page_data["text_regions"].append(region_data)
            
            data["pages"].append(page_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Exported {len(ocr_results)} pages to JSON")
        return output_path
    
    def _export_csv(self, ocr_results: List[OCROutput], output_path: Path) -> Path:
        """Export as CSV (one row per text region)"""
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            # Header
            writer.writerow([
                'Page', 'Region_ID', 'Text', 'Confidence',
                'X', 'Y', 'Width', 'Height'
            ])
            
            # Data
            for i, result in enumerate(ocr_results):
                for j, box in enumerate(result.text_boxes):
                    x, y, w, h = box.bbox
                    writer.writerow([
                        i + 1,
                        j + 1,
                        box.text,
                        f"{box.confidence:.3f}",
                        int(x),
                        int(y),
                        int(w),
                        int(h),
                    ])
        
        logger.info(f"Exported {len(ocr_results)} pages to CSV")
        return output_path
    
    def _export_docx(self, ocr_results: List[OCROutput], output_path: Path) -> Path:
        """Export as Microsoft Word document"""
        doc = DocxDocument()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('OCR Document', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Total Pages: {len(ocr_results)}")
        doc.add_paragraph("")
        
        # Add content
        for i, result in enumerate(ocr_results):
            if len(ocr_results) > 1:
                # Page header
                heading = doc.add_heading(f'Page {i+1}', level=2)
                
                # Add metadata
                meta = doc.add_paragraph()
                meta.add_run(f"Detected Regions: {result.num_detected_regions}  |  ")
                meta.add_run(f"Confidence: {result.confidence:.2%}")
                meta_format = meta.paragraph_format
                meta_format.space_after = Pt(6)
            
            # Add text content
            if result.full_text:
                para = doc.add_paragraph(result.full_text)
                para_format = para.paragraph_format
                para_format.space_after = Pt(12)
            else:
                para = doc.add_paragraph("[No text detected]")
                para.runs[0].italic = True
                para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Page break (except for last page)
            if i < len(ocr_results) - 1:
                doc.add_page_break()
        
        doc.save(str(output_path))
        logger.info(f"Exported {len(ocr_results)} pages to DOCX")
        return output_path
    
    def _export_pdf(
        self,
        ocr_results: List[OCROutput],
        output_path: Path,
        document: Optional[Document] = None
    ) -> Path:
        """Export as PDF (text-only)"""
        pdf_doc = fitz.open()
        
        for i, result in enumerate(ocr_results):
            # Create page
            page = pdf_doc.new_page(width=595, height=842)  # A4 size
            
            # Add title
            page.insert_text(
                (50, 50),
                f"Page {i+1}",
                fontsize=16,
                fontname="helv",
                color=(0, 0, 0),
            )
            
            # Add text content
            y_position = 80
            for line in result.full_text.split('\n'):
                if y_position > 800:  # Near bottom of page
                    break
                
                page.insert_text(
                    (50, y_position),
                    line,
                    fontsize=11,
                    fontname="helv",
                    color=(0, 0, 0),
                )
                y_position += 15
        
        pdf_doc.save(str(output_path))
        pdf_doc.close()
        
        logger.info(f"Exported {len(ocr_results)} pages to PDF")
        return output_path
    
    def _export_searchable_pdf(
        self,
        ocr_results: List[OCROutput],
        output_path: Path,
        document: Optional[Document] = None
    ) -> Path:
        """Export as searchable PDF (image + invisible text layer)"""
        if document is None:
            raise ExportException("Original document required for searchable PDF export")
        
        pdf_doc = fitz.open()
        
        for i, (page_data, ocr_result) in enumerate(zip(document.pages, ocr_results)):
            # Convert page image to bytes
            img = page_data.image
            is_success, buffer = cv2.imencode(".png", img)
            if not is_success:
                raise ExportException(f"Failed to encode page {i+1}")
            
            img_bytes = buffer.tobytes()
            
            # Create PDF page
            img_pdf = fitz.open(stream=img_bytes, filetype="png")
            page = pdf_doc.new_page(
                width=img_pdf[0].rect.width,
                height=img_pdf[0].rect.height,
            )
            
            # Insert image
            page.insert_image(page.rect, stream=img_bytes)
            
            # Add invisible text layer
            for box in ocr_result.text_boxes:
                if not box.text:
                    continue
                
                x, y, w, h = box.bbox
                
                # Scale coordinates if needed
                scale_x = page.rect.width / img.shape[1]
                scale_y = page.rect.height / img.shape[0]
                
                rect = fitz.Rect(
                    x * scale_x,
                    y * scale_y,
                    (x + w) * scale_x,
                    (y + h) * scale_y,
                )
                
                # Insert text (invisible)
                page.insert_textbox(
                    rect,
                    box.text,
                    fontsize=h * scale_y * 0.8,  # Approximate font size
                    color=(1, 1, 1),  # White text (invisible on white background)
                    fill=(1, 1, 1, 0),  # Transparent fill
                )
            
            img_pdf.close()
        
        pdf_doc.save(str(output_path))
        pdf_doc.close()
        
        logger.info(f"Exported {len(ocr_results)} pages to searchable PDF")
        return output_path
    
    def export_regions_overlay(
        self,
        image: np.ndarray,
        text_boxes: List[TextBox],
        output_path: Path,
    ) -> Path:
        """
        Export image with text region overlays (for visualization)
        
        Args:
            image: Source image
            text_boxes: Detected text boxes
            output_path: Output image path
        
        Returns:
            Path to output image
        """
        # Create copy
        overlay = image.copy()
        
        # Draw boxes
        for box in text_boxes:
            points = box.points.astype(np.int32)
            
            # Draw polygon
            cv2.polylines(overlay, [points], True, (0, 255, 0), 2)
            
            # Add text label
            x, y, _, h = box.bbox
            label = f"{box.text[:20]}... ({box.confidence:.2f})"
            
            # Background for text
            (text_w, text_h), _ = cv2.getTextSize(
                label, cv2.FONT_HERSHEY_SIMPLEX, 0.5, 1
            )
            cv2.rectangle(
                overlay,
                (x, y - text_h - 5),
                (x + text_w, y),
                (0, 255, 0),
                -1,
            )
            
            # Text
            cv2.putText(
                overlay,
                label,
                (x, y - 5),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.5,
                (0, 0, 0),
                1,
            )
        
        # Save
        cv2.imwrite(str(output_path), overlay)
        logger.info(f"Exported overlay image with {len(text_boxes)} regions")
        
        return output_path
